# generate_doc.py
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()

# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('博思AI智能体\n系统架构文档')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('版本：v3.0\n更新日期：2026年8月')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ========== 目录页 ==========
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概述', '二、技术栈总览', '三、系统架构',
    '四、后端模块详解', '五、前端模块详解', '六、大数据架构',
    '七、数据库设计', '八、API接口设计', '九、Redis设计',
    '十、消息队列设计', '十一、WebSocket与多节点通信设计',
    '十二、安全体系', '十三、多节点部署架构',
    '十四、性能与可观测性', '十五、扩展规划'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ========== 一、项目概述 ==========
doc.add_heading('一、项目概述', level=1)
doc.add_paragraph(
    '博思AI智能体是一个融合应用架构、中间件架构、基础设施与AI架构的企业级智能平台，'
    '整合了AI对话、观点辩论、知识图谱、多模态生成、AI多人游戏、大数据分析等核心能力，'
    '构建具备感知、规划、执行与反思能力的智能体系统。'
    '系统支持多节点水平扩展部署，通过RabbitMQ实现跨节点消息广播，保障分布式环境下的一致性体验。'
)

doc.add_heading('1.1 核心功能模块', level=2)
add_table(['模块', '说明'], [
    ['AI伙伴群聊', '多模型同时在线，群聊式AI对话体验'],
    ['观点辩论场', '三位AI专家并行辩论，综合观点生成深度结论'],
    ['个人对话空间', '私密对话空间，支持文件上传、多轮上下文继承'],
    ['知识脉络图', '3D神经网络可视化，展示知识关联与语义网络'],
    ['图片与视频', '文生图/文生视频，多模态内容一键生成'],
    ['AI多人游戏', 'AI乒乓球、AI蛇王争霸、AI城池攻防战等互动游戏'],
    ['问答足迹', '历史问答集中管理，支持检索与二次编辑'],
    ['SQL执行器', '在线数据库查询与执行工具'],
    ['在线人数监控', '实时在线人数曲线图，按天/小时粒度查看趋势'],
    ['大数据分析', 'Spark批处理 + Flink实时计算 + Hive元数据管理'],
])

# ========== 二、技术栈总览 ==========
doc.add_heading('二、技术栈总览', level=1)

doc.add_heading('2.1 前端技术栈', level=2)
add_table(['技术', '版本', '用途'], [
    ['React', '18.2.0', 'UI框架'],
    ['React Router', '6.x', 'SPA路由管理'],
    ['Vite', '5.x', '构建工具'],
    ['Axios', '1.4.0', 'HTTP客户端'],
    ['SockJS', '1.6.1', 'WebSocket降级方案'],
    ['STOMP.js', '7.0.0', '消息协议客户端'],
    ['Canvas 2D', '原生', '游戏引擎、图表绘制'],
])

doc.add_heading('2.2 后端技术栈', level=2)
add_table(['技术', '版本', '用途'], [
    ['Spring Boot', '3.1.6', '应用框架'],
    ['Spring Security', '6.1', '安全认证（Lambda DSL）'],
    ['Spring WebSocket', '内置', '实时通信（SimpleBroker模式）'],
    ['Spring AMQP', '内置', '跨节点消息桥接（RabbitMQ AMQP）'],
    ['MyBatis', '3.0.2', 'ORM持久化'],
    ['MySQL', '8.0', '关系型数据库'],
    ['Redis', '-', '缓存/会话/分布式锁/在线人数共享'],
    ['RabbitMQ', '-', '业务消息队列 + 跨节点消息总线'],
    ['JWT (jjwt)', '0.11.5', '身份认证'],
    ['Apache POI', '5.2.5', 'Office文件解析'],
    ['阿里云内容安全', 'green20220302', '敏感词/内容检测'],
])

doc.add_heading('2.3 大数据技术栈', level=2)
add_table(['技术', '用途', '说明'], [
    ['Apache Spark', '批处理/ETL/ML', '大规模数据清洗、转换、机器学习训练'],
    ['Apache Flink', '实时流计算/CEP', '实时事件处理、复杂事件模式匹配'],
    ['Hive MetaStore', '统一元数据管理', '跨引擎元数据共享，统一数据目录'],
    ['Airflow', '任务调度/工作流编排', 'DAG调度、任务依赖管理、重试机制'],
    ['ClickHouse', 'OLAP分析/实时查询', '列式存储，亚秒级聚合查询'],
    ['HDFS', '分布式文件存储', '海量数据持久化存储'],
    ['HBase', '列式海量数据', '十亿级行数据随机读写'],
    ['Elasticsearch', '全文检索/日志分析', '倒排索引，复杂搜索'],
    ['MinIO', '对象存储/文件管理', 'S3兼容，图片/视频/附件存储'],
])

doc.add_heading('2.4 AI模型接入', level=2)
add_table(['模型', '提供方', '能力'], [
    ['DeepSeek', '深度求索', '深度推理、逻辑分析'],
    ['豆包', '字节跳动', '联网搜索、通用对话'],
    ['千问 (qwen-plus/vl-max)', '阿里云', '多模态理解、整合结论'],
    ['qwen-image-2.0-pro', '阿里云DashScope', '文生图'],
    ['wan2.7-t2v', '通义万相', '文生视频（最长10秒）'],
])

# ========== 三、系统架构 ==========
doc.add_heading('三、系统架构', level=1)

doc.add_heading('3.1 整体架构分层（九层）', level=2)
layers = [
    ('前端层 · Frontend', 'React SPA · WebSocket · STOMP · Vite · React Router'),
    ('负载均衡层 · Load Balancer', 'Nginx反向代理 · ip_hash会话粘性 · WebSocket升级 · SSL终止'),
    ('网关层 · Gateway', 'Spring Boot · JWT鉴权 · 路由分发 · API限流'),
    ('安全层 · Security', 'Spring Security · 数据加密 · 日志审计 · 阿里云内容安全检测'),
    ('应用核心 · Application Core', 'Java 稳态层（用户/权限/消息/事务/游戏）'),
    ('AI服务层 · AI Services', 'DeepSeek · 豆包 · 千问 · 内容安全 · 图像生成 · 视频生成'),
    ('中间件层 · Middleware', 'RabbitMQ（业务消息 + 跨节点广播） · Redis（缓存/会话/锁）'),
    ('大数据层 · Big Data', 'Spark · Flink · Hive MetaStore · Airflow · ClickHouse · Elasticsearch'),
    ('存储层 · Storage', 'MySQL · Redis · HDFS · HBase · MinIO'),
]
add_table(['层级', '组件'], layers)

doc.add_heading('3.2 核心数据流', level=2)
doc.add_paragraph(
    '用户请求 → Nginx（ip_hash负载均衡/SSL终止/WebSocket升级）'
    '→ 前端React SPA（HTTP/WebSocket/STOMP）'
    '→ 应用核心(Java) / AI服务层 / 大数据层'
    '→ 中间件层(RabbitMQ业务消息 + 跨节点广播 / Redis缓存与会话共享)'
    '→ MySQL / Redis / HDFS / HBase / MinIO'
)

doc.add_heading('3.3 多节点消息流', level=2)
doc.add_paragraph(
    '节点A广播消息 → BroadcastService → 本地SimpleBroker推送 + RabbitMQ TopicExchange(cross-node)'
    '→ 节点B的CrossNodeMessageListener接收 → Base64解码 → 本地SimpleBroker推送'
    '→ 节点B的WebSocket客户端收到消息'
)

# ========== 四、后端模块详解 ==========
doc.add_heading('四、后端模块详解', level=1)

doc.add_heading('4.1 项目结构', level=2)
structure = [
    'src/main/java/com/example/chat/',
    '├── Application.java                    # 启动入口',
    '├── config/',
    '│   ├── CrossNodeConfig.java            # 跨节点通信配置（RabbitMQ交换机/队列/监听）',
    '│   ├── CrossNodeMessageListener.java   # 跨节点消息监听（Base64解码+本地广播）',
    '│   ├── GlobalExceptionHandler.java     # 全局异常处理',
    '│   ├── RabbitConfig.java               # RabbitMQ连接与交换机配置',
    '│   ├── SecurityConfig.java             # Spring Security配置',
    '│   ├── ViewConfig.java                 # SPA路由转发配置',
    '│   ├── WebSocketConfig.java            # WebSocket/STOMP配置（SimpleBroker）',
    '│   └── WebSocketSessionTracker.java    # WebSocket会话追踪（Redis共享）',
    '├── consumer/',
    '│   └── ChatRequestConsumer.java        # MQ消息消费者',
    '├── controller/',
    '│   ├── AuthController.java             # 登录/注册',
    '│   ├── MessageController.java          # 消息/对话',
    '│   ├── DebateController.java           # 观点辩论',
    '│   ├── MediaGenController.java         # 图片/视频生成',
    '│   ├── AttachmentController.java       # 文件上传',
    '│   ├── ModelConfigController.java      # 模型管理',
    '│   ├── MonitorController.java          # 在线人数监控',
    '│   ├── ProfileController.java          # 个人信息',
    '│   ├── SqlExecutorController.java      # SQL执行器',
    '│   ├── CastleSiegeBattlefieldController.java  # 城池攻防战战场',
    '│   ├── CastleSiegeLordController.java  # 城池攻防战领主',
    '│   └── WebPageController.java          # 页面路由',
    '├── entity/                             # 实体类',
    '├── repository/                         # MyBatis Mapper接口',
    '├── security/',
    '│   ├── JwtAuthenticationFilter.java',
    '│   └── JwtUtil.java',
    '└── service/',
    '    ├── BroadcastService.java           # 跨节点广播服务（本地+RabbitMQ）',
    '    ├── ChatProcessor.java              # 核心聊天处理',
    '    ├── DebateProcessor.java            # 辩论处理',
    '    ├── MessageService.java             # 消息服务',
    '    ├── ModelAutoChatService.java       # AI自动对话',
    '    ├── ContentSafetyService.java       # 内容安全检测',
    '    ├── RateLimitService.java           # 限流服务',
    '    ├── AuditService.java               # 审计服务',
    '    ├── OnlineCountRedisService.java    # 在线人数Redis服务',
    '    ├── OnlineCountScheduler.java       # 在线人数定时采集',
    '    ├── CastleSiegeBattlefieldService.java  # 城池攻防战战场服务',
    '    └── CastleSiegeLordService.java     # 城池攻防战领主服务',
]
for line in structure:
    p = doc.add_paragraph(line)
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

doc.add_heading('4.2 核心业务流程', level=2)

doc.add_heading('4.2.1 AI对话流程', level=3)
steps = [
    '前端发送问题 → POST /api/v1/messages',
    'Controller接收 → 校验参数 → 写入MySQL(status=queued)',
    '发布消息到RabbitMQ',
    'Consumer消费消息：检查Redis缓存 → 命中则直接返回 → 未命中则调用AI模型API',
    '回答通过BroadcastService广播 → 本地WebSocket推送 + RabbitMQ跨节点同步',
    '完整回答写入MySQL + Redis缓存',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('4.2.2 AI辩论流程', level=3)
debate_steps = [
    '用户提问',
    '3个模型并行辩论（3轮），每轮各模型可见全部历史讨论',
    '千问模型综合所有观点 → 生成整合结论',
    '组合原始问题 + 结论 → 千问生成最终回答',
    'BroadcastService广播 + MySQL持久化',
]
for i, s in enumerate(debate_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('4.2.3 多模态生成流程', level=3)
media_steps = [
    '用户输入提示词 → POST /api/v1/media/generate',
    '根据类型路由：图片 → qwen-image-2.0-pro，视频 → wan2.7-t2v',
    '异步调用 → 返回任务ID',
    '轮询查询结果 → 前端展示',
]
for i, s in enumerate(media_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

# ========== 五、前端模块详解 ==========
doc.add_heading('五、前端模块详解', level=1)

doc.add_heading('5.1 项目结构', level=2)
fe_structure = [
    'frontend/src/',
    '├── App.jsx                   # 路由 + 导航 + 在线追踪',
    '├── main.jsx                  # 入口文件',
    '├── pages/',
    '│   ├── Landing.jsx           # 首页',
    '│   ├── Chat.jsx              # AI伙伴群聊',
    '│   ├── Debate.jsx            # 观点辩论场',
    '│   ├── PersonalChat.jsx      # 个人对话空间',
    '│   ├── KnowledgeGraph.jsx    # 知识脉络图',
    '│   ├── MediaGen.jsx          # 图片与视频生成',
    '│   ├── History.jsx           # 问答足迹',
    '│   ├── Profile.jsx           # 个人信息',
    '│   ├── About.jsx             # 制作人简介',
    '│   ├── AdminModels.jsx       # 模型管理',
    '│   ├── SqlExecutor.jsx       # SQL执行器',
    '│   ├── Monitor.jsx           # 在线人数监控',
    '│   ├── pingpang.jsx          # AI乒乓球',
    '│   ├── snakeking.jsx         # AI蛇王争霸',
    '│   └── castlesiege.jsx       # AI城池攻防战',
    '└── styles/',
    '    ├── chat.css              # 聊天样式',
    '    ├── debate.css            # 辩论样式',
    '    ├── game.css              # 游戏独立样式',
    '    └── ...                   # 其他页面样式',
]
for line in fe_structure:
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

doc.add_heading('5.2 路由配置', level=2)
add_table(['路径', '页面', '说明'], [
    ['/home', 'Landing', '首页'],
    ['/', 'Chat', 'AI伙伴群聊'],
    ['/debate', 'Debate', '观点辩论场'],
    ['/personal', 'PersonalChat', '个人对话空间'],
    ['/graph', 'KnowledgeGraph', '知识脉络图'],
    ['/media', 'MediaGen', '图片与视频'],
    ['/history', 'History', '问答足迹'],
    ['/profile', 'Profile', '个人信息'],
    ['/about', 'About', '制作人简介'],
    ['/games/pingpong', 'PingPong', 'AI乒乓球'],
    ['/games/snakeking', 'SnakeKing', 'AI蛇王争霸'],
    ['/games/castlesiege', 'CastleSiege', 'AI城池攻防战'],
    ['/admin/models', 'AdminModels', '模型管理'],
    ['/sql', 'SqlExecutor', 'SQL执行器'],
    ['/monitor', 'Monitor', '在线监控'],
])

doc.add_heading('5.3 构建与部署', level=2)
add_table(['配置项', '值'], [
    ['构建工具', 'Vite 5.x'],
    ['基础路径', '/chat/'],
    ['构建命令', 'npm run build'],
    ['输出目录', 'src/main/resources/static/chat/'],
    ['开发代理', 'API → localhost:8081，WebSocket → localhost:8081'],
    ['SPA路由', '所有前端路由由后端ViewConfig转发到index.html'],
])

# ========== 六、大数据架构 ==========
doc.add_heading('六、大数据架构', level=1)

doc.add_heading('6.1 大数据层架构', level=2)
doc.add_paragraph(
    '数据采集层（Kafka/Flume）→ 计算层（Spark批处理 + Flink实时流计算）'
    '→ 元数据层（Hive MetaStore统一目录）→ 调度层（Airflow DAG编排）'
    '→ 存储/查询层（ClickHouse OLAP + HDFS分布式文件 + HBase列式存储）'
)

doc.add_heading('6.2 各组件职责', level=2)
add_table(['组件', '职责', '应用场景'], [
    ['Spark', '大规模批处理、ETL、ML', '日志分析、用户行为统计、推荐模型训练'],
    ['Flink', '实时流计算、CEP', '实时在线人数统计、异常行为检测、实时告警'],
    ['Hive MetaStore', '统一元数据管理', 'Spark/Flink共享元数据，统一数据治理'],
    ['Airflow', '工作流编排、DAG调度', '每日数据汇总、ETL流水线、定时报表'],
    ['ClickHouse', '列式OLAP分析', '在线人数趋势分析、用户活跃度报表'],
    ['HDFS', '分布式文件存储', '原始日志存储、模型训练数据集'],
    ['HBase', '列式NoSQL', '用户行为明细、消息历史归档'],
    ['Elasticsearch', '全文检索、日志分析', '问答内容搜索、审计日志查询'],
])

doc.add_heading('6.3 数据流转示例', level=2)
doc.add_heading('6.3.1 用户行为分析流水线', level=3)
doc.add_paragraph(
    '用户操作 → Kafka事件流 → Flink（实时在线人数统计 → Redis → 前端展示；'
    '异常行为检测 → 告警通知）+ HDFS（原始事件落盘）→ Airflow每日调度 → '
    'Spark ETL数据清洗 → Hive表 → Spark ML模型训练 → ClickHouse聚合分析 → 监控报表'
)

doc.add_heading('6.3.2 AI对话数据分析', level=3)
doc.add_paragraph(
    '对话记录 → MySQL + RabbitMQ → 实时：Flink统计对话量/响应时间 → 监控面板；'
    '离线：Spark分析高频问题/用户满意度 → Hive MetaStore统一元数据 → '
    'ClickHouse OLAP查询 → 运营报表'
)

# ========== 七、数据库设计 ==========
doc.add_heading('七、数据库设计', level=1)

doc.add_heading('7.1 核心表', level=2)
add_table(['表名', '说明', '关键字段'], [
    ['users', '用户表', 'id, email, password_hash, name, role'],
    ['messages', '消息表', 'id, req_id, user_id, question, answer(JSON), status, provider, model'],
    ['model_configs', '模型配置', 'id, provider, model, api_key_encrypted, priority, enabled'],
    ['attachments', '附件表', 'id, message_id, storage_url, mime_type, filename'],
    ['audit_logs', '审计日志', 'id, user_id, action, detail, ip, created_at'],
    ['debate_records', '辩论记录', 'id, question, rounds, models, conclusion, final_answer'],
    ['online_count_records', '在线人数', 'id, page, count, recorded_at'],
    ['user_registrations', '注册信息', 'id, guest_name, registered_at'],
])

doc.add_heading('7.2 MySQL DDL', level=2)
ddl_statements = [
    ('users', '''CREATE TABLE users (
  id BIGINT AUTO_INCREMENT PRIMARY KEY,
  email VARCHAR(255) NOT NULL UNIQUE,
  password_hash VARCHAR(255) NOT NULL,
  name VARCHAR(100),
  role VARCHAR(32) DEFAULT 'user',
  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
);'''),
    ('messages', '''CREATE TABLE messages (
  id BIGINT AUTO_INCREMENT PRIMARY KEY,
  req_id VARCHAR(64) NOT NULL UNIQUE,
  user_id BIGINT NOT NULL,
  question TEXT NOT NULL,
  answer JSON,
  status VARCHAR(32) DEFAULT 'queued',
  provider VARCHAR(64),
  model VARCHAR(128),
  tokens_used INT DEFAULT 0,
  latency_ms INT,
  is_private TINYINT(1) DEFAULT 0,
  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  INDEX idx_user_id (user_id),
  INDEX idx_status (status),
  INDEX idx_created (created_at)
);'''),
    ('model_configs', '''CREATE TABLE model_configs (
  id BIGINT AUTO_INCREMENT PRIMARY KEY,
  provider VARCHAR(64) NOT NULL,
  model VARCHAR(128) NOT NULL,
  api_key_encrypted TEXT,
  display_name VARCHAR(100),
  priority INT DEFAULT 0,
  enabled TINYINT(1) DEFAULT 1,
  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);'''),
    ('online_count_records', '''CREATE TABLE online_count_records (
  id BIGINT AUTO_INCREMENT PRIMARY KEY,
  page VARCHAR(64) NOT NULL,
  count INT NOT NULL DEFAULT 0,
  recorded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
  INDEX idx_page_time (page, recorded_at),
  INDEX idx_recorded (recorded_at)
);'''),
]
for name, ddl in ddl_statements:
    p = doc.add_paragraph(f'-- {name}')
    p.runs[0].bold = True
    p2 = doc.add_paragraph(ddl)
    for run in p2.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

# ========== 八、API接口设计 ==========
doc.add_heading('八、API接口设计', level=1)

doc.add_heading('8.1 认证接口', level=2)
add_table(['方法', '路径', '说明'], [
    ['POST', '/api/v1/auth/register', '用户注册'],
    ['POST', '/api/v1/auth/login', '用户登录，返回JWT'],
    ['GET', '/api/v1/auth/me', '获取当前用户信息'],
])

doc.add_heading('8.2 消息接口', level=2)
add_table(['方法', '路径', '说明'], [
    ['POST', '/api/v1/messages', '发送消息（触发AI回答）'],
    ['POST', '/api/v1/messages/with-file', '发送带文件的消息'],
    ['GET', '/api/v1/messages', '获取消息列表（分页）'],
    ['GET', '/api/v1/messages/{reqId}', '获取单条消息详情'],
    ['DELETE', '/api/v1/messages/{reqId}', '删除消息'],
])

doc.add_heading('8.3 模型管理接口', level=2)
add_table(['方法', '路径', '说明'], [
    ['GET', '/api/v1/models', '获取模型配置列表'],
    ['POST', '/api/v1/models', '新增模型配置'],
    ['PUT', '/api/v1/models/{id}', '更新模型配置'],
    ['DELETE', '/api/v1/models/{id}', '删除模型配置'],
])

doc.add_heading('8.4 多模态生成接口', level=2)
add_table(['方法', '路径', '说明'], [
    ['POST', '/api/v1/media/generate', '生成图片/视频'],
    ['GET', '/api/v1/media/{taskId}', '查询生成任务状态'],
])

doc.add_heading('8.5 其他接口', level=2)
add_table(['方法', '路径', '说明'], [
    ['POST', '/api/v1/attachments/upload', '文件上传'],
    ['GET', '/api/v1/monitor/online', '在线人数监控数据'],
    ['POST', '/api/v1/sql/execute', '执行SQL查询'],
    ['GET', '/api/v1/profile', '获取个人信息'],
    ['PUT', '/api/v1/profile', '更新个人信息'],
])

# ========== 九、Redis设计 ==========
doc.add_heading('九、Redis设计', level=1)
add_table(['Key模式', '用途', 'TTL'], [
    ['chat:cache:{questionHash}:{model}', 'AI回答缓存', '24小时'],
    ['ws:page:{pageKey}', 'WebSocket会话集合（多节点共享）', '持久'],
    ['ws:known:pages', '已知页面集合', '持久'],
    ['ws:visit:{page}:{date}', '日访问次数统计', '持久'],
    ['online:count:{page}', '实时在线人数', '持久'],
    ['rate:{userId}:{minute}', 'API限流计数', '1分钟'],
    ['lock:debate:{questionHash}', '辩论分布式锁', '30秒'],
    ['media:task:{taskId}', '多模态生成任务状态', '1小时'],
])

# ========== 十、消息队列设计 ==========
doc.add_heading('十、消息队列设计', level=1)

doc.add_heading('10.1 业务消息队列', level=2)
add_table(['类型', '名称', '绑定', '说明'], [
    ['Direct', 'chat.request', 'chat.request.queue', 'AI对话请求'],
    ['Direct', 'chat.response', 'chat.response.queue', 'AI对话响应'],
    ['Direct', 'media.generate', 'media.generate.queue', '多模态生成任务'],
    ['Topic', 'audit.events', 'audit.log.queue', '审计日志事件'],
    ['Delayed', 'chat.retry', 'chat.retry.queue', '失败重试（延迟队列）'],
])

doc.add_heading('10.2 跨节点消息总线', level=2)
add_table(['类型', '名称', '路由键', '说明'], [
    ['Topic', 'cross-node', '#（全路由）', '跨节点WebSocket消息广播'],
])
doc.add_paragraph(
    '每个节点启动时自动创建独立队列（cross-node-{nodeId}），绑定到cross-node交换机。'
    'BroadcastService发送消息时，同时通过本地SimpleBroker推送和RabbitMQ广播。'
    'CrossNodeMessageListener监听本节点队列，收到其他节点消息后在本地重新广播，实现跨节点消息同步。'
    '通过消息体中的_nodeId字段防止回环（跳过自己发出的消息）。'
)

# ========== 十一、WebSocket与多节点通信设计 ==========
doc.add_heading('十一、WebSocket与多节点通信设计', level=1)

doc.add_heading('11.1 WebSocket配置', level=2)
add_table(['配置项', '值', '说明'], [
    ['Broker模式', 'SimpleBroker', '内存级单节点代理，不依赖外部STOMP代理'],
    ['STOMP端点', '/ws/chat', 'WebSocket连接入口（SockJS降级）'],
    ['订阅前缀', '/topic', '客户端订阅路径前缀'],
    ['发送前缀', '/app', '客户端发送路径前缀'],
    ['心跳间隔', '25秒', '双向心跳保活'],
    ['会话超时', '300秒', 'WebSocket空闲超时'],
    ['消息大小限制', '128KB', '单条消息最大体积'],
])

doc.add_heading('11.2 STOMP目的地', level=2)
add_table(['STOMP目的地', '方向', '说明'], [
    ['/topic/user.{userId}', '服务端→客户端', '用户私有消息推送（AI回答/错误）'],
    ['/topic/public-questions', '服务端→客户端', '公共对话广播（群聊/自动对话）'],
    ['/topic/online-count/{page}', '服务端→客户端', '页面在线人数广播'],
    ['/topic/online-count/all', '服务端→客户端', '全站在线人数广播'],
    ['/topic/debate.{userId}', '服务端→客户端', '辩论进度推送'],
    ['/topic/castlesiege.state', '服务端→客户端', '城池攻防战状态同步'],
])

doc.add_heading('11.3 跨节点通信架构', level=2)
add_table(['组件', '职责'], [
    ['BroadcastService', '统一广播入口：本地SimpleBroker推送 + RabbitMQ跨节点发布'],
    ['CrossNodeConfig', 'RabbitMQ TopicExchange(cross-node) + 每节点独立Queue + 监听容器'],
    ['CrossNodeMessageListener', '监听本节点队列，Base64解码后本地广播，跳过自身消息防回环'],
    ['WebSocketSessionTracker', '通过Redis共享会话信息，通过BroadcastService同步在线人数'],
])

doc.add_heading('11.4 跨节点消息流时序', level=2)
cross_node_steps = [
    '节点A调用 broadcastService.broadcast("/topic/online-count/landing", data)',
    'BroadcastService执行 messagingTemplate.convertAndSend() → 节点A本地WebSocket客户端收到消息',
    'BroadcastService执行 rabbitTemplate.convertAndSend("cross-node", destination, payload)',
    'RabbitMQ将消息路由到所有绑定队列（节点A队列 + 节点B队列）',
    '节点A的CrossNodeMessageListener收到消息，检测_nodeId等于自身 → 跳过（防回环）',
    '节点B的CrossNodeMessageListener收到消息，检测_nodeId不等于自身 → 执行本地广播',
    '节点B的WebSocket客户端通过SimpleBroker收到消息',
]
for i, s in enumerate(cross_node_steps, 1):
    doc.add_paragraph(f'{i}. {s}')

# ========== 十二、安全体系 ==========
doc.add_heading('十二、安全体系', level=1)

doc.add_heading('12.1 认证与授权', level=2)
doc.add_paragraph('• JWT Token认证：登录后签发，有效期24小时，支持刷新')
doc.add_paragraph('• Spring Security Lambda DSL配置，前后端分离架构')
doc.add_paragraph('• 接口级权限控制：管理员接口需ADMIN角色')

doc.add_heading('12.2 内容安全', level=2)
doc.add_paragraph('• 阿里云内容安全SDK（green20220302）')
doc.add_paragraph('• 用户输入敏感词检测，AI输出内容审核')
doc.add_paragraph('• 违规内容自动拦截并记录审计日志')

doc.add_heading('12.3 数据安全', level=2)
doc.add_paragraph('• API密钥AES加密存储')
doc.add_paragraph('• 密码BCrypt哈希存储')
doc.add_paragraph('• SQL注入防护（MyBatis参数化查询）')
doc.add_paragraph('• XSS防护（Spring Security默认转义）')

# ========== 十三、多节点部署架构 ==========
doc.add_heading('十三、多节点部署架构', level=1)

doc.add_heading('13.1 部署拓扑', level=2)
doc.add_paragraph(
    'Nginx（反向代理/ip_hash负载均衡/SSL终止）→ 多节点Spring Boot应用（8081 + 8082）'
    '→ MySQL 8.0 + Redis + RabbitMQ（共享中间件）'
)

doc.add_heading('13.2 Nginx负载均衡配置', level=2)
add_table(['配置项', '值', '说明'], [
    ['upstream', 'backend_nodes', '后端节点池'],
    ['负载策略', 'ip_hash', '基于客户端IP的会话粘性'],
    ['节点1', '127.0.0.1:8081', '主节点'],
    ['节点2', '127.0.0.1:8082', '从节点'],
    ['WebSocket升级', 'Upgrade + Connection头', '支持HTTP→WebSocket协议升级'],
    ['超时时间', '3600秒', 'WebSocket长连接保持'],
    ['配置路径', '/etc/nginx/conf.d/chat-system.conf', 'Nginx配置文件'],
])

doc.add_heading('13.3 多节点数据一致性', level=2)
add_table(['数据类型', '共享方式', '说明'], [
    ['WebSocket会话', 'Redis（ws:page:*）', '所有节点共享会话-页面映射'],
    ['在线人数', 'Redis + BroadcastService', 'Redis统计 + 跨节点实时同步'],
    ['AI回答缓存', 'Redis', '所有节点共享同一缓存'],
    ['业务消息', 'RabbitMQ', '所有节点竞争消费'],
    ['广播消息', 'RabbitMQ TopicExchange', '所有节点副本接收'],
    ['用户数据', 'MySQL', '所有节点共享同一数据库'],
    ['API限流', 'Redis', '全局限流，不区分节点'],
])

doc.add_heading('13.4 构建与部署流程', level=2)
add_table(['步骤', '命令', '说明'], [
    ['1. 前端构建', 'cd frontend && npm run build', 'Vite构建输出到static/chat/'],
    ['2. 后端打包', 'mvn clean package -DskipTests', '生成可执行JAR'],
    ['3. 上传JAR', 'scp target/*.jar root@server:/opt/app/', '上传到服务器'],
    ['4. 启动节点1', 'java -jar app.jar --server.port=8081', '启动主节点'],
    ['5. 启动节点2', 'java -jar app.jar --server.port=8082', '启动从节点'],
    ['6. Nginx重载', 'nginx -t && nginx -s reload', '重载负载均衡配置'],
])

doc.add_heading('13.5 节点扩缩容', level=2)
doc.add_paragraph(
    '扩容：启动新节点（--server.port=8083），Nginx upstream添加新server即可。'
    '每个节点启动时自动创建独立RabbitMQ队列并绑定到cross-node交换机，自动加入消息广播网络。'
    '缩容：停止节点进程，RabbitMQ队列自动清理（auto-delete），Nginx移除对应server。'
)

# ========== 十四、性能与可观测性 ==========
doc.add_heading('十四、性能与可观测性', level=1)
add_table(['维度', '方案', '说明'], [
    ['缓存', 'Redis多级缓存', 'AI回答缓存、会话缓存、在线人数'],
    ['异步', 'RabbitMQ消息队列', 'AI请求异步处理，削峰填谷'],
    ['限流', 'Redis滑动窗口', '每用户每分钟请求上限'],
    ['日志', 'SLF4J + Logback', '结构化日志，按级别分文件'],
    ['监控', '在线人数Redis+MySQL', '实时采集，按天/小时聚合查询'],
    ['审计', 'AuditService + audit_logs表', '全操作审计，IP记录'],
    ['负载均衡', 'Nginx ip_hash', '会话粘性，均匀分布'],
    ['水平扩展', 'RabbitMQ跨节点广播', '无状态节点，按需扩容'],
])

# ========== 十五、扩展规划 ==========
doc.add_heading('十五、扩展规划', level=1)
add_table(['方向', '规划', '优先级'], [
    ['大数据落地', '部署Spark/Flink集群，接入实时分析', '高'],
    ['知识图谱增强', 'Neo4j图数据库，自动知识抽取', '高'],
    ['游戏扩展', '更多AI游戏类型，排行榜系统', '中'],
    ['微服务拆分', '核心服务拆分为独立微服务', '中'],
    ['容器化部署', 'Docker Compose / K8s编排', '中'],
    ['多租户支持', '租户隔离、资源配额管理', '低'],
    ['国际化', 'i18n多语言支持', '低'],
])

# ========== 保存 ==========
output_path = '/Users/apple/IdeaProjects/chat-system-project/docs/博思AI智能体-系统架构文档.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
